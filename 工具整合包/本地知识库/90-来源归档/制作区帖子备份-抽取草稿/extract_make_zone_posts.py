#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build reusable index drafts for the local CCZ mod forum backup.

The script intentionally does not copy whole forum posts into official topic
documents. It extracts enough metadata and short candidates to route each
source into the existing knowledge base.
"""

from __future__ import annotations

import csv
import hashlib
import html
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Iterable


SCRIPT_DIR = Path(__file__).resolve().parent
KB_ROOT = SCRIPT_DIR.parent.parent
SOURCE_ROOT = KB_ROOT / "制作区帖子备份"

OUTPUT_SOURCES = SCRIPT_DIR / "sources.csv"
OUTPUT_TOPIC_MAP = SCRIPT_DIR / "topic-map.md"
OUTPUT_SUMMARY = SCRIPT_DIR / "extract-summary.md"
OUTPUT_UNRESOLVED = SCRIPT_DIR / "unresolved.md"

IN_SCOPE_SUFFIXES = {".htm", ".html", ".doc", ".docx", ".txt"}
IGNORED_SUFFIXES = {".rar"}
MAX_SUMMARY_TEXT = 120_000


TOPIC_RULES = [
    {
        "topic": "EXE逆向与函数地址",
        "keywords": [
            "exe",
            "olly",
            "ollydbg",
            "koei",
            "代码",
            "函数",
            "汇编",
            "攻击响应",
            "伤害",
            "反击",
            "地址",
        ],
        "targets": [
            "04-函数速查/函数速查表.md",
            "01-核心引擎/引擎工作索引.md",
            "05-教程指南/战场操作自动化与调试流程.md",
        ],
    },
    {
        "topic": "star175版本与整型变量",
        "keywords": [
            "star175",
            "star5",
            "5.6",
            "5.7",
            "5.8",
            "5.9",
            "6.0",
            "6.1",
            "整形变量",
            "整型变量",
            "发布",
            "正式版",
        ],
        "targets": [
            "09-版本与外部资料/6X版本差异与适配.md",
            "02-数据结构/全局变量.md",
            "08-剧本与战场/全剧本指令详解2023归纳.md",
        ],
    },
    {
        "topic": "godtype扩展引擎",
        "keywords": [
            "godtype",
            "最终引擎",
            "data扩展",
            "扩展引擎",
            "旧引擎",
            "新引擎",
            "指针变量",
        ],
        "targets": [
            "01-核心引擎/引擎工作索引.md",
            "02-数据结构/全局变量.md",
            "03-机制详解/特效值机制.md",
        ],
    },
    {
        "topic": "兵种宝物特效机制",
        "keywords": [
            "兵种",
            "宝物",
            "特效",
            "致命一击",
            "mp恢复",
            "hp回复",
            "策略",
            "地形",
            "相克",
            "恢复mp",
        ],
        "targets": [
            "03-机制详解/特效值机制.md",
            "07-数据表与资源/兵种设定.md",
            "07-数据表与资源/宝物物品设定.md",
        ],
    },
    {
        "topic": "剧本R/S变量指令",
        "keywords": [
            "剧本",
            "指令",
            "整型",
            "变量",
            "r单挑",
            "单挑",
            "71特效请求",
            "75特殊形象",
            "信息传送",
        ],
        "targets": [
            "08-剧本与战场/全剧本指令详解2023归纳.md",
            "08-剧本与战场/剧本指令大全/README.md",
            "02-数据结构/全局变量.md",
        ],
    },
    {
        "topic": "资源美工形象音效",
        "keywords": [
            "美工",
            "形象",
            "换色",
            "mcall",
            "meff",
            "头像",
            "半身头像",
            "音效",
            "对话框",
            "颜色化",
        ],
        "targets": [
            "07-数据表与资源/策略动画-Meff-Mcall.md",
            "07-数据表与资源/人物形象-RS-形象指定器.md",
            "07-数据表与资源/图片处理-E5资源.md",
        ],
    },
    {
        "topic": "制作教程与工具流程",
        "keywords": [
            "教程",
            "实际操作",
            "做自己的",
            "修改详解",
            "操作演示",
            "使用ollydbg",
        ],
        "targets": [
            "05-教程指南/入门指南.md",
            "05-教程指南/战场操作自动化与调试流程.md",
            "06-项目与工具链/验证与烟测.md",
        ],
    },
]


@dataclass
class SourceRecord:
    rel_path: str
    suffix: str
    size: int
    mtime: str
    sha256: str
    title: str = ""
    author_candidates: list[str] = field(default_factory=list)
    date_candidates: list[str] = field(default_factory=list)
    topics: list[str] = field(default_factory=list)
    targets: list[str] = field(default_factory=list)
    version_scope: str = "未知旧论坛资料"
    evidence_level: str = "外部佐证/待证"
    status: str = "待处理"
    notes: str = ""
    addresses: list[str] = field(default_factory=list)
    variables: list[str] = field(default_factory=list)
    resources: list[str] = field(default_factory=list)
    text_sample: str = ""


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def kb_rel(path: Path) -> str:
    return path.relative_to(KB_ROOT).as_posix()


def source_rel(path: Path) -> str:
    return path.relative_to(SOURCE_ROOT).as_posix()


def decode_bytes(data: bytes, preferred: Iterable[str] = ()) -> tuple[str, str]:
    if data.startswith(b"\xef\xbb\xbf"):
        return data.decode("utf-8-sig", errors="replace"), "utf-8-sig"
    if data.startswith(b"\xff\xfe"):
        return data.decode("utf-16le", errors="replace"), "utf-16le"
    if data.startswith(b"\xfe\xff"):
        return data.decode("utf-16be", errors="replace"), "utf-16be"

    encodings = list(preferred) + ["utf-8", "gb18030", "gbk", "big5"]
    seen: set[str] = set()
    for enc in encodings:
        normalized = enc.lower()
        if normalized in seen:
            continue
        seen.add(normalized)
        try:
            return data.decode(enc), enc
        except UnicodeDecodeError:
            continue
    return data.decode("gb18030", errors="replace"), "gb18030-replace"


def detect_html_charset(data: bytes) -> str | None:
    head = data[:4096].decode("ascii", errors="ignore")
    match = re.search(r"charset\s*=\s*['\"]?([A-Za-z0-9_-]+)", head, re.I)
    return match.group(1) if match else None


def clean_text(text: str) -> str:
    text = html.unescape(text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = []
    for line in text.split("\n"):
        line = re.sub(r"\s+", " ", line).strip()
        if not line:
            continue
        lines.append(line)
    return "\n".join(lines)


def extract_html(path: Path) -> tuple[str, str, str, list[str], list[str]]:
    data = path.read_bytes()
    charset = detect_html_charset(data) or "gb18030"
    raw, used_encoding = decode_bytes(data, [charset])

    title = ""
    title_match = re.search(r"<title[^>]*>(.*?)</title>", raw, re.I | re.S)
    if title_match:
        title = clean_text(re.sub(r"<[^>]+>", " ", title_match.group(1)))

    body = re.sub(r"<script[\s\S]*?</script>", " ", raw, flags=re.I)
    body = re.sub(r"<style[\s\S]*?</style>", " ", body, flags=re.I)
    body = re.sub(r"</(?:p|div|br|tr|td|li|h[1-6]|table)>", "\n", body, flags=re.I)
    body = re.sub(r"<[^>]+>", " ", body)
    text = clean_text(body)

    dates = sorted(set(re.findall(r"发表于\s*([0-9]{4}-[0-9]{1,2}-[0-9]{1,2}\s+[0-9]{1,2}:[0-9]{2})", text)))
    authors = []
    for match in re.finditer(r"(?:报告\s+发红包\s+回复\s+|只看该作者\s+)([^\n#]{2,24})\s+#\d+", text):
        candidate = match.group(1).strip()
        if candidate and candidate not in authors:
            authors.append(candidate)
        if len(authors) >= 5:
            break

    return title, text, used_encoding, authors, dates[:5]


def extract_docx(path: Path) -> tuple[str, str]:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - environment dependent
        return "", f"python-docx unavailable: {exc}"

    doc = Document(path)
    chunks: list[str] = []
    for para in doc.paragraphs:
        value = para.text.strip()
        if value:
            chunks.append(value)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    text = clean_text("\n".join(chunks))
    if text:
        return text, ""

    # Some old docx files store visible text in shapes/text boxes not exposed by
    # python-docx. Fall back to raw OOXML text nodes before declaring failure.
    try:
        with zipfile.ZipFile(path) as archive:
            xml_names = [name for name in archive.namelist() if name.startswith("word/") and name.endswith(".xml")]
            xml_texts: list[str] = []
            for name in xml_names:
                raw = archive.read(name).decode("utf-8", errors="replace")
                for match in re.finditer(r"<w:t[^>]*>(.*?)</w:t>", raw, flags=re.S):
                    xml_texts.append(html.unescape(re.sub(r"<[^>]+>", "", match.group(1))))
            text = clean_text("\n".join(xml_texts))
            if text:
                return text, "raw OOXML fallback"
    except Exception as exc:
        return "", f"python-docx empty; raw OOXML fallback failed: {exc}"

    return "", "python-docx and raw OOXML found no visible text"


def try_convert_doc(path: Path) -> tuple[str, str]:
    """Try available command-line converters for legacy .doc files."""
    antiword = shutil.which("antiword")
    if antiword:
        try:
            proc = subprocess.run(
                [antiword, str(path)],
                check=False,
                capture_output=True,
                timeout=60,
            )
            if proc.returncode == 0 and proc.stdout.strip():
                text, _ = decode_bytes(proc.stdout, ["gb18030"])
                return clean_text(text), "antiword"
        except Exception:
            pass

    office = shutil.which("soffice") or shutil.which("libreoffice")
    if office:
        with tempfile.TemporaryDirectory() as tmp:
            out_dir = Path(tmp)
            try:
                proc = subprocess.run(
                    [
                        office,
                        "--headless",
                        "--convert-to",
                        "txt:Text",
                        "--outdir",
                        str(out_dir),
                        str(path),
                    ],
                    check=False,
                    capture_output=True,
                    timeout=90,
                )
                if proc.returncode == 0:
                    txt_files = list(out_dir.glob("*.txt"))
                    if txt_files:
                        text, _ = decode_bytes(txt_files[0].read_bytes(), ["utf-8", "gb18030"])
                        return clean_text(text), "libreoffice"
            except Exception:
                pass

    return "", "legacy .doc converter unavailable"


def infer_version(rel_path: str, title: str) -> str:
    hay = f"{rel_path} {title}".lower()
    if "godtype" in hay:
        return "godtype旧扩展引擎资料"
    if "star175" in hay or "star" in hay:
        versions = []
        for version in ["5.6", "5.7", "5.8", "5.9", "6.0", "6.1"]:
            if version in hay:
                versions.append(version)
        return "star175 " + ("/".join(versions) if versions else "旧版资料")
    if "koei" in hay or "原版" in hay:
        return "KOEI原版/旧EXE逆向资料"
    if "6.5" in hay:
        return "6.5相关外部资料，需本地复核"
    if re.search(r"\b5\.[6-9]\b|\b6\.[01]\b", hay):
        return "5.x/6.0/6.1旧版资料"
    return "未知旧论坛资料"


def add_topic(topics: list[str], targets: list[str], topic: str) -> None:
    for rule in TOPIC_RULES:
        if rule["topic"] != topic:
            continue
        if topic not in topics:
            topics.append(topic)
        for target in rule["targets"]:
            if target not in targets:
                targets.append(target)
        return


def classify(rel_path: str, title: str, text: str) -> tuple[list[str], list[str]]:
    # Forum pages contain a lot of navigation and related-thread text, so route
    # primarily from path/title and only use body text for high-signal evidence.
    name_hay = f"{rel_path}\n{title}".lower()
    body_hay = text[:MAX_SUMMARY_TEXT].lower()
    topics: list[str] = []
    targets: list[str] = []
    address_count = len(set(re.findall(r"\b0?0?4[0-9a-f]{5}\b", body_hay)))

    if (
        any(key in name_hay for key in ["exe", "olly", "koei", "代码", "函数", "补充研究", "修改详解"])
        or address_count >= 8
    ):
        add_topic(topics, targets, "EXE逆向与函数地址")

    if (
        "star175" in name_hay
        or re.search(r"\bstar5|\bstar6|5\.[6-9]|6\.[01]|正式版|修正版|发一个exe", name_hay)
        or "整形变量" in name_hay
    ):
        add_topic(topics, targets, "star175版本与整型变量")

    if "godtype" in name_hay or "最终引擎" in name_hay or "data扩展" in name_hay or "旧引擎" in name_hay:
        add_topic(topics, targets, "godtype扩展引擎")

    if any(key in name_hay for key in ["兵种", "宝物", "特效", "策略", "地形", "相克", "致命一击", "hp回复", "mp"]):
        add_topic(topics, targets, "兵种宝物特效机制")

    if any(key in name_hay for key in ["剧本", "指令", "整型", "变量", "r单挑", "单挑", "71特效请求", "75特殊形象", "章节"]):
        add_topic(topics, targets, "剧本R/S变量指令")

    if any(key in name_hay for key in ["美工", "形象", "换色", "mcall", "meff", "头像", "音效", "对话框", "颜色化"]):
        add_topic(topics, targets, "资源美工形象音效")

    if any(key in name_hay for key in ["教程", "实际操作", "做自己的", "操作演示", "使用ollydbg"]):
        add_topic(topics, targets, "制作教程与工具流程")

    if not topics:
        # Body-only classification stays conservative and requires repeated
        # evidence so ordinary forum chrome does not overroute the source.
        if body_hay.count("mcall") >= 2 or body_hay.count("meff") >= 2:
            add_topic(topics, targets, "资源美工形象音效")
        if body_hay.count("整型变量") >= 2 or body_hay.count("信息传送") >= 2:
            add_topic(topics, targets, "剧本R/S变量指令")
        if body_hay.count("宝物特效") >= 2 or body_hay.count("兵种") >= 4:
            add_topic(topics, targets, "兵种宝物特效机制")

    if not topics:
        topics = ["综合旧帖待分流"]
        targets = ["90-来源归档/散落资料整合索引.md"]
    return topics, targets


def collect_tokens(text: str) -> tuple[list[str], list[str], list[str]]:
    addresses = sorted(set(re.findall(r"\b0?0?4[0-9A-Fa-f]{5}\b", text.upper())))[:20]
    variables = sorted(set(re.findall(r"\b40[0-9]{2}\b", text)))[:20]
    resources = sorted(
        set(
            re.findall(
                r"\b(?:[A-Za-z_][A-Za-z0-9_]*\.e5|Mcall[0-9A-Za-z_-]*\.e5|Meff\.e5|Ekd5\.exe|Data\.e5|Star\.e5)\b",
                text,
                flags=re.I,
            )
        )
    )[:20]
    return addresses, variables, resources


def make_record(path: Path) -> SourceRecord:
    stat = path.stat()
    rel = source_rel(path)
    suffix = path.suffix.lower()
    record = SourceRecord(
        rel_path=f"制作区帖子备份/{rel}",
        suffix=suffix,
        size=stat.st_size,
        mtime=datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
        sha256=sha256_file(path),
    )

    text = ""
    if suffix in {".htm", ".html"}:
        title, text, encoding, authors, dates = extract_html(path)
        record.title = title
        record.author_candidates = authors
        record.date_candidates = dates
        record.status = "已抽取HTML摘要"
        record.notes = f"HTML解码={encoding}；正文仅用于摘要和分类，未写入全文"
    elif suffix == ".docx":
        text, err = extract_docx(path)
        record.title = path.stem
        record.status = "已抽取DOCX摘要" if text else "DOCX抽取失败"
        record.notes = err if err else "python-docx抽取段落和表格文本"
    elif suffix == ".doc":
        text, converter = try_convert_doc(path)
        record.title = path.stem
        if text:
            record.status = "已抽取DOC摘要"
            record.notes = f"legacy .doc converted by {converter}"
        else:
            record.status = "待抽取"
            record.notes = converter
    elif suffix == ".txt":
        record.title = path.stem
        if stat.st_size == 0:
            record.status = "空来源/无内容"
            record.notes = "空txt，仅登记文件存在"
        else:
            text, encoding = decode_bytes(path.read_bytes(), ["utf-8", "gb18030", "gbk"])
            text = clean_text(text)
            record.status = "已抽取TXT摘要"
            record.notes = f"TXT解码={encoding}"

    if not record.title:
        record.title = path.stem

    record.addresses, record.variables, record.resources = collect_tokens(text)
    record.version_scope = infer_version(record.rel_path, record.title)
    record.topics, record.targets = classify(record.rel_path, record.title, text)
    if suffix == ".docx" and not text and not record.notes:
        record.notes = "python-docx未抽取到段落或表格文本，可能是图片型或特殊docx"
    record.text_sample = text[:MAX_SUMMARY_TEXT]
    return record


def normalize_title(title: str) -> str:
    title = re.sub(r"\s+", "", title)
    title = re.sub(r"[-_ ]?轩辕春秋文化论坛.*$", "", title)
    title = re.sub(r"\(?[0-9]+\)?$", "", title)
    title = title.replace("(2)", "").replace("2-", "")
    return title


def build_summary(record: SourceRecord) -> list[str]:
    bullets = [
        f"主题候选：{'、'.join(record.topics)}。",
        f"建议归入：{'、'.join(record.targets[:4])}。",
        f"版本边界：{record.version_scope}；证据等级保持为 `{record.evidence_level}`。",
        "复用方式：作为阅读导航、候选地址/变量/指令/资源线索和后续验证队列，不直接作为 6.5 写回规则。",
    ]
    if record.addresses:
        bullets.append("地址候选：" + "、".join(f"`{item}`" for item in record.addresses[:10]) + "。")
    if record.variables:
        bullets.append("整型变量候选：" + "、".join(f"`{item}`" for item in record.variables[:10]) + "。")
    if record.resources:
        bullets.append("资源/文件名候选：" + "、".join(f"`{item}`" for item in record.resources[:10]) + "。")
    if record.status.startswith("待") or record.status.startswith("空") or "失败" in record.status:
        bullets.append(f"处理限制：{record.status}；{record.notes}。")
    else:
        bullets.append(f"抽取状态：{record.status}；{record.notes}。")
    return bullets[:8]


def write_sources(records: list[SourceRecord]) -> None:
    fields = [
        "rel_path",
        "type",
        "size",
        "mtime",
        "sha256",
        "title",
        "author_candidates",
        "date_candidates",
        "version_scope",
        "topics",
        "target_docs",
        "evidence_level",
        "status",
        "notes",
    ]
    with OUTPUT_SOURCES.open("w", encoding="utf-8-sig", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=fields)
        writer.writeheader()
        for rec in records:
            writer.writerow(
                {
                    "rel_path": rec.rel_path,
                    "type": rec.suffix.lstrip("."),
                    "size": rec.size,
                    "mtime": rec.mtime,
                    "sha256": rec.sha256,
                    "title": rec.title,
                    "author_candidates": ";".join(rec.author_candidates),
                    "date_candidates": ";".join(rec.date_candidates),
                    "version_scope": rec.version_scope,
                    "topics": ";".join(rec.topics),
                    "target_docs": ";".join(rec.targets),
                    "evidence_level": rec.evidence_level,
                    "status": rec.status,
                    "notes": rec.notes,
                }
            )


def write_topic_map(records: list[SourceRecord]) -> None:
    groups: dict[str, list[SourceRecord]] = {}
    for rec in records:
        for topic in rec.topics:
            groups.setdefault(topic, []).append(rec)

    lines = [
        "# 制作区帖子备份主题映射",
        "",
        "## 结论速览",
        "",
        "- 本页由 `extract_make_zone_posts.py` 自动生成，记录制作区帖子备份的主题归口。",
        "- 所有来源默认是外部佐证、旧口径或待证线索；正式写入前必须回到 6.5 本地样本、旧工具源码、复读、烟测或实机验证。",
        "- `star175.rar` 按本轮范围要求忽略，未进入本主题映射。",
        "",
        "## 主题归口",
        "",
        "| 主题 | 数量 | 建议主入口 | 来源 |",
        "|------|------|------------|------|",
    ]
    for topic in sorted(groups):
        recs = sorted(groups[topic], key=lambda item: item.rel_path)
        targets = []
        for rec in recs:
            for target in rec.targets:
                if target not in targets:
                    targets.append(target)
        source_list = "<br>".join(f"`{rec.rel_path}`" for rec in recs[:12])
        if len(recs) > 12:
            source_list += f"<br>...另 {len(recs) - 12} 个"
        lines.append(f"| {topic} | {len(recs)} | {'<br>'.join(f'`{target}`' for target in targets[:5])} | {source_list} |")

    OUTPUT_TOPIC_MAP.write_text("\n".join(lines) + "\n", encoding="utf-8-sig")


def write_summary(records: list[SourceRecord]) -> None:
    lines = [
        "# 制作区帖子备份抽取摘要",
        "",
        "## 结论速览",
        "",
        "- 本页是自动摘要草稿，不是正式知识条目。",
        "- 摘要只保留主题、候选地址/变量/资源和验证边界，不复制论坛长帖全文。",
        "- `.doc` 如缺少转换器会标为待抽取；不阻塞 HTML、DOCX 和 TXT 的主题映射。",
        "",
        "## 来源摘要",
        "",
    ]
    for rec in sorted(records, key=lambda item: item.rel_path):
        lines.append(f"### {rec.rel_path}")
        lines.append("")
        lines.append(f"- 标题：{rec.title}")
        for bullet in build_summary(rec):
            lines.append(f"- {bullet}")
        lines.append("")

    OUTPUT_SUMMARY.write_text("\n".join(lines), encoding="utf-8-sig")


def write_unresolved(records: list[SourceRecord], ignored: list[Path]) -> None:
    duplicate_groups: dict[str, list[SourceRecord]] = {}
    for rec in records:
        duplicate_groups.setdefault(normalize_title(rec.title), []).append(rec)
    duplicates = {k: v for k, v in duplicate_groups.items() if k and len(v) > 1}

    unresolved = [
        rec
        for rec in records
        if rec.status.startswith("待") or rec.status.startswith("空") or "失败" in rec.status
    ]
    lines = [
        "# 制作区帖子备份未决项",
        "",
        "## 结论速览",
        "",
        "- 本页记录自动抽取阶段无法直接升级为正式知识的项。",
        "- `.doc` 老格式在当前环境缺少转换器时只登记来源；后续可补 Word、LibreOffice 或 antiword 后重跑脚本。",
        "- `star175.rar` 按用户本轮范围要求忽略。",
        "",
        "## 忽略项",
        "",
    ]
    if ignored:
        for path in ignored:
            lines.append(f"- `{kb_rel(path)}`：本轮明确忽略，不进入 sources.csv。")
    else:
        lines.append("- 无。")

    lines.extend(["", "## 待抽取或无内容", ""])
    if unresolved:
        for rec in unresolved:
            lines.append(f"- `{rec.rel_path}`：{rec.status}；{rec.notes}")
    else:
        lines.append("- 无。")

    lines.extend(["", "## 疑似分页或重复主题", ""])
    if duplicates:
        for _, recs in sorted(duplicates.items()):
            lines.append("- " + "；".join(f"`{rec.rel_path}`" for rec in sorted(recs, key=lambda item: item.rel_path)))
    else:
        lines.append("- 无。")

    lines.extend(
        [
            "",
            "## 后续验证队列",
            "",
            "- 地址类来源进入 `Ekd5.exe` 静态复读和 x32dbg 动态命中队列。",
            "- 剧本类来源进入 `.eex` 样本复读、旧剧本编辑器源码对照和实机流程验证队列。",
            "- 资源类来源进入 E5/DLL 资源读取、编号映射和预览复核队列。",
            "- 特效类来源进入当前 6.5 特效号、效果值、触发链和烟测候选队列。",
        ]
    )
    OUTPUT_UNRESOLVED.write_text("\n".join(lines) + "\n", encoding="utf-8-sig")


def main() -> int:
    if not SOURCE_ROOT.exists():
        print(f"source root not found: {SOURCE_ROOT}", file=sys.stderr)
        return 2

    records: list[SourceRecord] = []
    ignored: list[Path] = []
    for path in sorted(SOURCE_ROOT.rglob("*")):
        if not path.is_file():
            continue
        suffix = path.suffix.lower()
        if suffix in IGNORED_SUFFIXES:
            ignored.append(path)
            continue
        if suffix not in IN_SCOPE_SUFFIXES:
            continue
        records.append(make_record(path))

    write_sources(records)
    write_topic_map(records)
    write_summary(records)
    write_unresolved(records, ignored)

    counts: dict[str, int] = {}
    for rec in records:
        counts[rec.suffix] = counts.get(rec.suffix, 0) + 1
    count_text = ", ".join(f"{key or '(none)'}={counts[key]}" for key in sorted(counts))
    print(f"MAKE_ZONE_POSTS_EXTRACT_OK records={len(records)} {count_text} ignored={len(ignored)}")
    print("outputs=sources.csv; topic-map.md; extract-summary.md; unresolved.md")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
